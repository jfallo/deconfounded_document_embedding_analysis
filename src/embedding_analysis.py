import numpy as np
import scipy as sp
import pandas as pd
import sklearn
import sklearn.metrics
import matplotlib.pyplot as plt
from collections import Counter
from docx import Document

np.set_printoptions(precision= 4, suppress= True)

# load .csv data and store text column in a numpy array
df = pd.read_csv('input/wiki_lexis_data.csv', quotechar='"')
df['id'] = df.index
wiki_data = df[df.values[:,1] == 'wiki'].values[:,0]
wiki_idxs = np.array(df[df.values[:,1] == 'wiki'].values[:,2], dtype= int)
lexis_data = df[df.values[:,1] == 'lexis'].values[:,0]
lexis_idxs = np.array(df[df.values[:,1] == 'lexis'].values[:,2], dtype= int)

# load embeddings .npy data and store contents in respective numpy arrays
original_embeddings = np.load('input/wiki_lexis_embeddings.npy')
LEACE_embeddings = np.load('input/wiki_lexis_embeddings_erased.npy')

wiki_original_embeddings = original_embeddings[wiki_idxs]
wiki_LEACE_embeddings = LEACE_embeddings[wiki_idxs]

lexis_original_embeddings = original_embeddings[lexis_idxs]
lexis_LEACE_embeddings = LEACE_embeddings[lexis_idxs]


# compute cosine similarity matrices
wiki_original_similarity = sklearn.metrics.pairwise.cosine_similarity(wiki_original_embeddings)
wiki_LEACE_similarity = sklearn.metrics.pairwise.cosine_similarity(wiki_LEACE_embeddings)

lexis_original_similarity = sklearn.metrics.pairwise.cosine_similarity(lexis_original_embeddings)
lexis_LEACE_similarity = sklearn.metrics.pairwise.cosine_similarity(lexis_LEACE_embeddings)


def get_significant_pairs(S, T, delta= 99.99):
    """
    Input: 
        S: original similarity matrix
        T: LEACE similarity matrix
        delta: percentile
    Output: pairs that see a significant decrease / increase in similarity after LEACE
    """
    similarity_dec = np.maximum(S - T, 0)
    pairs_significantly_less_similar = np.argwhere(similarity_dec > np.percentile(similarity_dec, delta))
    
    similarity_inc = np.maximum(T - S, 0)
    pairs_significantly_more_similar = np.argwhere(similarity_inc > np.percentile(similarity_inc, delta))

    return np.array(pairs_significantly_less_similar, dtype= int), np.array(pairs_significantly_more_similar, dtype= int)


def plot_pairs(pairs, shape, title, file_name):
    I = np.zeros(shape)

    for pair in pairs:
        I[pair[0], pair[1]] = 1

    Y, X = np.where(I == 1)

    plt.scatter(X, Y, c= 'blue', s= 2)
    plt.grid(True)
    plt.title(title)
    plt.savefig(f'output/{file_name}')


def write_texts(text_idxs, data, file_name):
    doc = Document()

    for i in text_idxs:
        p = doc.add_paragraph(f'Text {i}: ', style= 'List Number')
        p.add_run(data[i])

    doc.save(f'output/{file_name}.docx')


K = 20


# === WIKI ANALYSIS ===
pairs_significantly_less_similar, pairs_significantly_more_similar = get_significant_pairs(wiki_original_similarity, wiki_LEACE_similarity)


# === pairs significantly less similar analysis ===
# K texts that appear most often in significantly less similar pairs
top_K_significantly_less_similar = np.array(Counter(pairs_significantly_less_similar[:,0]).most_common())[:K,0]
write_texts(top_K_significantly_less_similar, wiki_data, 'wiki_significantly_less_similar')

# text pairs that are significantly less similar
unique_pairs_significantly_less_similar = {(min(i, j), max(i, j)) for (i, j) in pairs_significantly_less_similar}
unique_pairs_significantly_less_similar = list(unique_pairs_significantly_less_similar)
plot_pairs(
    unique_pairs_significantly_less_similar, 
    wiki_original_similarity.shape, 
    'Pairs Significantly Less Similar (wiki)',
    'wiki_significantly_less_similar_pairs'
)


# === pairs significantly more similar analysis ===
# K texts that appear most often in significantly more similar pairs
top_K_significantly_more_similar = np.array(Counter(pairs_significantly_more_similar[:,0]).most_common())[:K,0]
write_texts(top_K_significantly_more_similar, wiki_data, 'wiki_significantly_more_similar')

# text pairs that are significantly more similar
unique_pairs_significantly_more_similar = {(min(i, j), max(i, j)) for (i, j) in pairs_significantly_more_similar}
unique_pairs_significantly_more_similar = list(unique_pairs_significantly_more_similar)
plot_pairs(
    unique_pairs_significantly_more_similar, 
    wiki_original_similarity.shape, 
    'Pairs Significantly More Similar (wiki)',
    'wiki_significantly_more_similar_pairs'
)


# === LEXIS ANALYSIS ===
pairs_significantly_less_similar, pairs_significantly_more_similar = get_significant_pairs(lexis_original_similarity, lexis_LEACE_similarity)


# === pairs significantly less similar analysis ===
# K texts that appear most often in significantly less similar pairs
top_K_significantly_less_similar = np.array(Counter(pairs_significantly_less_similar[:,0]).most_common())[:K,0]
write_texts(top_K_significantly_less_similar, lexis_data, 'lexis_significantly_less_similar')

# text pairs that are significantly less similar
unique_pairs_significantly_less_similar = {(min(i, j), max(i, j)) for (i, j) in pairs_significantly_less_similar}
unique_pairs_significantly_less_similar = list(unique_pairs_significantly_less_similar)
plot_pairs(
    unique_pairs_significantly_less_similar, 
    lexis_original_similarity.shape, 
    'Pairs Significantly Less Similar (lexis)',
    'lexis_significantly_less_similar_pairs'
)


# === pairs significantly more similar analysis ===
# K texts that appear most often in significantly more similar pairs
top_K_significantly_more_similar = np.array(Counter(pairs_significantly_more_similar[:,0]).most_common())[:K,0]
write_texts(top_K_significantly_more_similar, lexis_data, 'lexis_significantly_more_similar')

# text pairs that are significantly more similar
unique_pairs_significantly_more_similar = {(min(i, j), max(i, j)) for (i, j) in pairs_significantly_more_similar}
unique_pairs_significantly_more_similar = list(unique_pairs_significantly_more_similar)
plot_pairs(
    unique_pairs_significantly_more_similar, 
    lexis_original_similarity.shape, 
    'Pairs Significantly More Similar (lexis)',
    'lexis_significantly_more_similar_pairs'
)




def get_K_nearest(S, data, K= 2):
    """
    Input: 
        S: similarity matrix
        data: text data
        K: some positive integer
    Output: the K most similar texts to text i where most similar texts to i are stored in row i
    """
    K_nearest = []
    nearest = np.argsort(S)

    for i in range(S.shape[0]):
        K_nearest.append(np.flip(data[nearest[i, -(K+1):-1]]))

    return np.array(K_nearest)


def get_K_furthest(S, data, K= 2):
    """
    Input: 
        S: similarity matrix
        data: text data
        K: some positive integer
    Output: the K most similar texts to text i where least similar texts to i are stored in row i
    """
    K_nearest = []
    nearest = np.argsort(S)

    for i in range(S.shape[0]):
        K_nearest.append(data[nearest[i, :K]])

    return np.array(K_nearest)


def write_K_nearest(S, data, source, K= 2):
    """
    Generates and save document of K most similar texts to text i where text i is bolded and followed by enumerated list of K most similar texts.
    Input: 
        S: similarity matrix
        data: text data
        source: 'wiki' or 'lexis'
        K: some positive integer
    """
    K_nearest = get_K_nearest(S, data, K)
    doc = Document()

    for i, section in enumerate(K_nearest):
        doc.add_heading(f'Text {i}:', level= 1)
        p = doc.add_paragraph()
        p.add_run(data[i]).bold = True
        
        for text in section:
            doc.add_paragraph(text, style= 'List Number')

    doc.save(f'output/{source}_{K}_nearest.docx')


def write_K_furthest(S, data, source, K= 2):
    """
    Generates and save document of K least similar texts to text i where text i is bolded and followed by enumerated list of K least similar texts.
    Input: 
        S: similarity matrix
        data: text data
        source: 'wiki' or 'lexis'
        K: some positive integer
    """
    K_nearest = get_K_furthest(S, data, K)
    doc = Document()

    for i, section in enumerate(K_nearest):
        doc.add_heading(f'Text {i}:', level= 1)
        p = doc.add_paragraph()
        p.add_run(data[i]).bold = True
        
        for text in section:
            doc.add_paragraph(text, style= 'List Number')

    doc.save(f'output/{source}_{K}_furthest.docx')


def plt_consistency(C, metric, source, K):
    """
    Generates and saves a bar plot of consistency scores.
    Input: 
        C: consistency matrix
        metrix: 'nearest' or 'furthest'
        source: 'wiki' or 'lexis'
        K: some positive integer
    """
    labels = np.round(np.linspace(0, 1, K+1), 4)
    counts = np.zeros(K+1)

    for i in range(K+1):
        counts[i] = np.sum(C == labels[i])

    plt.bar(np.array(labels, dtype= str), counts)
    plt.xlabel('Consistency Score')
    plt.ylabel('Number of Texts')
    plt.title(f'Original to LEACE consistency scores for {K} {metric} texts ({source})')
    plt.savefig(f'output/{source}_{K}_{metric}_consistency')


def compare_K_nearest(S, T, data, source, K= 2, plot= False):
    """
    Gives consistency score to each texts K nearest texts from S to T. Generates and saves results in a bar plot if plot= True.
    ex: Suppose K=10 and consider text i. 
        If 9 of 10 nearest texts for S and T are the same, then the consistency score of text i is 0.9 from S to T
    Input: 
        S: original similarity matrix
        T: LEACE similarity matrix
        data: text data
        source: 'wiki' or 'lexis'
        K: some positive integer
        plot: boolean
    Output: consistency score array
    """
    S_K_nearest = get_K_nearest(S, data, K)
    T_K_nearest = get_K_nearest(T, data, K)

    C = np.zeros(S.shape[0])

    for i in range(S.shape[0]):
        for text in S_K_nearest[i]:
            if text in T_K_nearest[i]:
                C[i] += 1
    
    C /= K

    if plot:
        plt_consistency(C, 'nearest', source, K)

    return C


def compare_K_furthest(S, T, data, source, K= 2, plot= False):
    """
    Gives consistency score to each texts K furthest texts from S to T. Generates and saves results in a bar plot if plot= True.
    ex: Suppose K=10 and consider text i. 
        If 9 of 10 furthest texts for S and T are the same, then the consistency score of text i is 0.9 from S to T
    Input: 
        S: original similarity matrix
        T: LEACE similarity matrix
        data: text data
        source: 'wiki' or 'lexis'
        K: some positive integer
        plot: boolean
    Output: consistency score array
    """
    S_K_furthest = get_K_furthest(S, data, K)
    T_K_furthest = get_K_furthest(T, data, K)

    C = np.zeros(S.shape[0])

    for i in range(S.shape[0]):
        for text in S_K_furthest[i]:
            if text in T_K_furthest[i]:
                C[i] += 1
    
    C /= K

    if plot:
        plt_consistency(C, 'furthest', source, K)

    return C
